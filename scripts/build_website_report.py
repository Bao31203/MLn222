from __future__ import annotations

from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Bao-cao-gioi-thieu-website-MLN222.docx"
SKILL_SCRIPTS = Path(
    r"C:\Users\pgb31\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.601.10930\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


# narrative_proposal preset with one named MLN222 brand-color override.
PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 100, "bottom": 100, "start": 120, "end": 120}

FONT = "Calibri"
INK = RGBColor(0x14, 0x26, 0x1F)
MUTED = RGBColor(0x5C, 0x6E, 0x66)
GREEN = RGBColor(0x1D, 0x5E, 0x4A)
GREEN_2 = RGBColor(0x2D, 0x74, 0x5F)
GOLD = RGBColor(0xB4, 0x70, 0x00)
GOLD_DARK = RGBColor(0x7C, 0x52, 0x00)
PALE_GREEN = "EAF3EE"
PALE_GOLD = "FFF5DC"
PALE_GRAY = "F4F6F5"
BORDER = "B8C7C0"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

STUDY_SHOT = (
    ROOT
    / "plans"
    / "260719-1752-hybrid-ui-command-wheel"
    / "reports"
    / "phase-02-study-shell-1440x900.png"
)
FLASHCARD_SHOT = (
    ROOT
    / "plans"
    / "260719-1752-hybrid-ui-command-wheel"
    / "reports"
    / "phase-03-flashcard-revealed-1440x900.png"
)
MAP_SHOT = (
    ROOT
    / "plans"
    / "260719-1752-hybrid-ui-command-wheel"
    / "reports"
    / "phase-04-tactical-map-1440x900.png"
)
WHEEL_SHOT = (
    ROOT
    / "plans"
    / "260719-1752-hybrid-ui-command-wheel"
    / "reports"
    / "phase-05-command-wheel-1440x900.png"
)
QUIZ_SHOT = (
    ROOT
    / "plans"
    / "260719-1752-hybrid-ui-command-wheel"
    / "reports"
    / "phase-08-quiz-review-390x844.png"
)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    name: str = FONT,
):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_style_font(style, size: float, color: RGBColor, bold: bool = False):
    style.font.name = FONT
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, spec in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in spec.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_table_borders(table, color: str = BORDER, size: int = 5):
    spec = {"val": "single", "sz": size, "space": 0, "color": color}
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(
                cell,
                top=spec,
                left=spec,
                bottom=spec,
                right=spec,
                insideH=spec,
                insideV=spec,
            )


def remove_table_borders(table):
    spec = {"val": "nil", "sz": 0, "space": 0, "color": "auto"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, top=spec, left=spec, bottom=spec, right=spec)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_field(run, instruction: str, fallback: str = ""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1D5E4A")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), FONT)
    run_pr.extend([fonts, color, underline])
    run.append(run_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_numbering(doc: Document, *, number_format: str, level_text: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), number_format)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), level_text)
    justify = OxmlElement("w:lvlJc")
    justify.set(qn("w:val"), "left")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "279")
    p_pr.extend([tabs, indent])
    level.extend([start, fmt, text, justify, suffix, p_pr])
    abstract.append(level)
    # Numbering schema requires every abstractNum before every concrete num.
    # Keeping that order prevents Word from repairing bullet definitions into decimals.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def add_list_item(doc, num_id: int, text: str, *, lead: str | None = None):
    p = doc.add_paragraph(style="List Paragraph")
    apply_numbering(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if lead:
        set_run_font(p.add_run(lead), bold=True, color=GREEN)
        set_run_font(p.add_run(" " + text.lstrip()), color=INK)
    else:
        set_run_font(p.add_run(text.lstrip()), color=INK)
    return p


def add_body(doc, text: str, *, bold_lead: str | None = None, align=None):
    p = doc.add_paragraph(style="Normal")
    if align is not None:
        p.alignment = align
    if bold_lead:
        set_run_font(p.add_run(bold_lead), bold=True, color=GREEN)
    set_run_font(p.add_run(text), color=INK)
    return p


def add_heading(doc, text: str, level: int, *, page_break_before: bool = False):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    if page_break_before:
        p.paragraph_format.page_break_before = True
    return p


def add_callout(doc, label: str, text: str, *, fill: str = PALE_GREEN, accent: str = "1D5E4A"):
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(
        table,
        [PAGE_WIDTH_DXA],
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa={"top": 140, "bottom": 140, "start": 190, "end": 190},
    )
    remove_table_borders(table)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_borders(
        cell,
        left={"val": "single", "sz": 18, "space": 0, "color": accent},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(f"{label}: "), bold=True, color=GREEN)
    set_run_font(p.add_run(text), color=INK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=1, cols=len(metrics))
    widths = [PAGE_WIDTH_DXA // len(metrics)] * len(metrics)
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa={"top": 150, "bottom": 150, "start": 90, "end": 90},
    )
    remove_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, (number, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, PALE_GOLD if idx % 2 else PALE_GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(str(number)), size=17, bold=True, color=GREEN)
        p.add_run().add_break(WD_BREAK.LINE)
        set_run_font(p.add_run(label), size=9.2, color=MUTED)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=9.3, align=None):
    p = cell.paragraphs[0]
    clear_paragraph(p)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    set_run_font(p.add_run(str(text)), size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_data_table(doc, headers, rows, widths, *, center_columns=(), total_row=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[idx],
            header,
            bold=True,
            color=WHITE,
            size=9.2,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shade_cell(table.rows[0].cells[idx], "1D5E4A")
    set_repeat_table_header(table.rows[0])

    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[col_idx], value, align=align)
            if row_idx % 2:
                shade_cell(cells[col_idx], PALE_GRAY)

    if total_row is not None:
        cells = table.add_row().cells
        for col_idx, value in enumerate(total_row):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[col_idx], value, bold=True, color=GREEN, align=align)
            shade_cell(cells[col_idx], PALE_GREEN)

    apply_table_geometry(
        table,
        widths,
        table_width_dxa=PAGE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS_DXA,
    )
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def set_image_alt(inline_shape, alt_text: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text[:120])


def add_figure(doc, path: Path, caption: str, *, width: float, alt_text: str):
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_image_alt(shape, alt_text)
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.keep_with_next = False
    c.paragraph_format.keep_together = True
    set_run_font(c.add_run(caption), size=9.5, italic=True, color=MUTED)
    return shape


def add_source_entry(doc, num_id: int, label: str, value: str, url: str | None = None):
    p = doc.add_paragraph(style="List Paragraph")
    apply_numbering(p, num_id)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(label), bold=True, color=GREEN)
    if url:
        add_hyperlink(p, value, url)
    else:
        set_run_font(p.add_run(value), color=INK)
    return p


def add_toc_field(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    set_run_font(run, color=INK)
    add_field(run, 'TOC \\o "1-1" \\h \\z \\u', "Đang tạo mục lục...")


def configure_styles(doc: Document):
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, 11, INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, 30, GREEN, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)
    title_r_pr = title._element.get_or_add_rPr()
    for tag in ("w:spacing", "w:kern"):
        node = title_r_pr.find(qn(tag))
        if node is not None:
            title_r_pr.remove(node)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, 15, GREEN_2)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p_pr = subtitle._element.get_or_add_pPr()
    subtitle_num_pr = subtitle_p_pr.find(qn("w:numPr"))
    if subtitle_num_pr is not None:
        subtitle_p_pr.remove(subtitle_num_pr)
    subtitle_r_pr = subtitle._element.get_or_add_rPr()
    subtitle_spacing = subtitle_r_pr.find(qn("w:spacing"))
    if subtitle_spacing is not None:
        subtitle_r_pr.remove(subtitle_spacing)

    h1 = styles["Heading 1"]
    set_style_font(h1, 16, GREEN, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = styles["Heading 2"]
    set_style_font(h2, 13, GREEN_2, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = styles["Heading 3"]
    set_style_font(h3, 12, GOLD_DARK, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    set_style_font(caption, 9.5, MUTED)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_with_next = False
    caption.paragraph_format.keep_together = True

    list_style = styles["List Paragraph"]
    set_style_font(list_style, 11, INK)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.208
    list_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name, size, indent in (("TOC 1", 9.5, 0), ("TOC 2", 9.0, 0.18)):
        if name in styles:
            style = styles[name]
            set_style_font(style, size, GREEN if name == "TOC 1" else INK, bold=name == "TOC 1")
            style.paragraph_format.left_indent = Inches(indent)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(1)
            style.paragraph_format.line_spacing = 1.0

    if "Cover Meta" not in styles:
        cover_meta = styles.add_style("Cover Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cover_meta = styles["Cover Meta"]
    set_style_font(cover_meta, 10, MUTED)
    cover_meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_meta.paragraph_format.space_after = Pt(5)
    cover_meta.paragraph_format.line_spacing = 1.0


def configure_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("MLN222"), size=8.5, bold=True, color=GREEN)
    set_run_font(p.add_run("\tBÁO CÁO GIỚI THIỆU WEBSITE"), size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("Trang "), size=8.5, color=MUTED)
    page_run = p.add_run()
    set_run_font(page_run, size=8.5, color=MUTED)
    add_field(page_run, "PAGE", "1")
    set_run_font(p.add_run(" / "), size=8.5, color=MUTED)
    total_run = p.add_run()
    set_run_font(total_run, size=8.5, color=MUTED)
    add_field(total_run, "NUMPAGES", "1")

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def build_report():
    for image_path in (STUDY_SHOT, FLASHCARD_SHOT, MAP_SHOT, WHEEL_SHOT, QUIZ_SHOT):
        if not image_path.exists():
            raise FileNotFoundError(image_path)

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    bullet_num = create_numbering(doc, number_format="bullet", level_text="•")
    decimal_num = create_numbering(doc, number_format="decimal", level_text="%1.")
    source_num = create_numbering(doc, number_format="decimal", level_text="%1.")

    props = doc.core_properties
    props.title = "Báo cáo giới thiệu website học tập MLN222"
    props.subject = "Giới thiệu sản phẩm, chức năng, kỹ thuật và kiểm thử"
    props.author = "MLN222 Quiz"
    props.keywords = "MLN222, Kinh tế chính trị Mác-Lênin, quiz, flashcard, Công thành"
    props.comments = "Báo cáo được tạo từ phiên bản website ngày 19/07/2026."

    # Cover: editorial_cover pattern, adapted to the site's visual identity.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    set_run_font(kicker.add_run("BÁO CÁO GIỚI THIỆU SẢN PHẨM"), size=10.5, bold=True, color=GOLD)

    title = doc.add_paragraph(style="Title")
    title.add_run("WEBSITE HỌC TẬP MLN222")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Ôn tập Kinh tế chính trị Mác-Lênin\nkết hợp game chiến lược Công thành")

    meta = doc.add_paragraph(style="Cover Meta")
    meta.add_run("Phiên bản 1.0  |  Tháng 7 năm 2026")
    add_figure(
        doc,
        MAP_SHOT,
        "Giao diện chiến dịch Công thành trên bản đồ Việt Nam 34 tỉnh/thành.",
        width=6.18,
        alt_text="Ảnh giao diện game Công thành với bản đồ Việt Nam và bảng điều hành chiến dịch.",
    )
    cover_note = doc.add_paragraph(style="Cover Meta")
    cover_note.paragraph_format.space_before = Pt(4)
    set_run_font(
        cover_note.add_run("Tài liệu mô tả mục tiêu, chức năng, trải nghiệm sử dụng và nền tảng kỹ thuật của website."),
        size=9.5,
        color=MUTED,
    )
    doc.add_page_break()

    add_heading(doc, "Tóm tắt điều hành", 1)
    add_callout(
        doc,
        "Tổng quan",
        "MLN222 Quiz là website tự học chạy trực tiếp trong trình duyệt, kết hợp luyện câu hỏi có dẫn nguồn với game chiến lược theo lượt. Sản phẩm hướng tới việc biến quá trình ôn tập lý thuyết thành một chu trình chủ động: học, kiểm tra, nhận phản hồi, áp dụng kiến thức và duy trì động lực qua tiến trình chiến dịch.",
    )
    add_metric_strip(
        doc,
        [
            ("504", "câu hỏi có giải thích"),
            ("6", "chương học phần"),
            ("34", "tỉnh/thành trong game"),
            ("4", "chế độ sử dụng"),
        ],
    )
    add_heading(doc, "Giá trị nổi bật", 2)
    add_list_item(doc, bullet_num, " Học tập có căn cứ: mỗi câu hỏi gồm bốn phương án, một đáp án, phần giải thích và dẫn nguồn theo trang giáo trình.", lead="Nội dung:")
    add_list_item(doc, bullet_num, " Luyện thi, Flashcard và Tìm kiếm phục vụ ba nhu cầu khác nhau nhưng dùng chung một ngân hàng kiến thức.", lead="Đa phương thức:")
    add_list_item(doc, bullet_num, " Game Công thành biến kết quả học thành tài nguyên và hiệu ứng, qua đó gắn kiến thức với quyết định chiến lược.", lead="Học qua chơi:")
    add_list_item(doc, bullet_num, " Toàn bộ ứng dụng được đóng gói trong một file HTML, chạy ngoại tuyến và tự động lưu tiến độ trên trình duyệt.", lead="Thuận tiện:")

    add_heading(doc, "Đối tượng sử dụng", 2)
    add_body(
        doc,
        "Website phù hợp với sinh viên học môn MLN222, người cần ôn tập Kinh tế chính trị Mác-Lênin theo từng chương, và giảng viên muốn có một công cụ luyện tập bổ trợ. Người học có thể dùng nhanh như một ngân hàng trắc nghiệm hoặc tham gia chiến dịch dài hạn để duy trì việc ôn tập qua nhiều lượt chơi.",
    )
    doc.add_page_break()

    toc_heading = doc.add_paragraph()
    toc_heading.paragraph_format.space_after = Pt(14)
    set_run_font(toc_heading.add_run("Mục lục"), size=20, bold=True, color=GREEN)
    add_toc_field(doc)
    add_heading(doc, "1. Tổng quan về website", 1, page_break_before=True)
    add_heading(doc, "1.1. Bối cảnh và định hướng", 2)
    add_body(
        doc,
        "Khối lượng lý thuyết của học phần MLN222 trải rộng từ lịch sử các học thuyết kinh tế, hàng hóa và thị trường, giá trị thặng dư, độc quyền, kinh tế thị trường định hướng xã hội chủ nghĩa đến công nghiệp hóa và hội nhập kinh tế quốc tế. Việc chỉ đọc lại tài liệu dễ tạo cảm giác quen thuộc nhưng chưa chắc giúp người học nhớ chính xác khái niệm, quan hệ và điều kiện áp dụng.",
    )
    add_body(
        doc,
        "Website được xây dựng theo hướng thực hành truy hồi kiến thức. Câu hỏi, phản hồi tức thời và giải thích giúp người học phát hiện điểm yếu; chế độ Flashcard hỗ trợ ôn nhanh; công cụ tìm kiếm giúp tra cứu; còn game Công thành tạo một vòng lặp dài hạn để kết quả học tập có tác động trực tiếp đến diễn biến chơi.",
    )

    add_heading(doc, "1.2. Mục tiêu sản phẩm", 2)
    add_list_item(doc, bullet_num, " Cung cấp bộ câu hỏi bao quát đủ sáu chương, có mức độ từ nhận biết đến vận dụng.")
    add_list_item(doc, bullet_num, " Phản hồi rõ ràng ngay sau mỗi lựa chọn, kèm giải thích và nguồn để người học tự kiểm chứng.")
    add_list_item(doc, bullet_num, " Duy trì hứng thú bằng cơ chế tài nguyên, lãnh thổ, ngoại giao và chiến đấu theo lượt.")
    add_list_item(doc, bullet_num, " Hoạt động ổn định trên máy tính và điện thoại, kể cả khi không có kết nối Internet.")
    add_list_item(doc, bullet_num, " Lưu trạng thái học và chiến dịch để người dùng có thể tiếp tục ở lần mở sau.")

    add_heading(doc, "1.3. Cấu trúc trải nghiệm", 2)
    add_body(
        doc,
        "Ứng dụng có bốn mục chính: Luyện thi, Flashcard, Tìm kiếm và Công thành. Ba mục học tập dùng giao diện sáng, ưu tiên khả năng đọc; game dùng giao diện tối dạng bản đồ chiến thuật. Trên máy tính, người dùng chuyển chế độ bằng thanh điều hướng bên trái; trên thiết bị di động, thanh điều hướng được đặt ở đáy màn hình.",
    )
    add_figure(
        doc,
        STUDY_SHOT,
        "Hình 1. Không gian Luyện thi với bộ lọc, thống kê phiên và câu hỏi trắc nghiệm.",
        width=5.0,
        alt_text="Giao diện Luyện thi MLN222 trên máy tính, gồm bộ lọc và câu hỏi bốn lựa chọn.",
    )

    add_heading(doc, "2. Nguồn học liệu và ngân hàng câu hỏi", 1)
    add_heading(doc, "2.1. Nguồn biên soạn", 2)
    add_body(
        doc,
        "Nguồn chuẩn của ngân hàng là giáo trình Kinh tế chính trị Mác-Lênin trong thư mục F:\\MLN222. Các tệp bài giảng theo từng slot được dùng để nhận diện trọng tâm và ví dụ bổ trợ. Khi cách diễn đạt giữa giáo trình và slide khác nhau, đáp án được xác định theo giáo trình PDF.",
    )
    add_callout(
        doc,
        "Nguyên tắc học liệu",
        "Không dùng số liệu thời sự hoặc ví dụ ngoài giáo trình làm căn cứ duy nhất cho đáp án. Câu hỏi phải có một đáp án rõ ràng, phương án nhiễu cùng phạm trù và nguồn dẫn đủ cụ thể để kiểm tra lại.",
        fill=PALE_GOLD,
        accent="B47000",
    )

    add_heading(doc, "2.2. Quy mô và phân bố", 2)
    question_rows = [
        ("1", "Khái quát môn học, đối tượng, phương pháp và chức năng", "64", "26 / 26 / 12"),
        ("2", "Hàng hóa, tiền tệ, thị trường và các quy luật kinh tế", "89", "36 / 36 / 17"),
        ("3", "Giá trị thặng dư, tích lũy và các hình thức biểu hiện", "99", "40 / 40 / 19"),
        ("4", "Cạnh tranh, độc quyền và chủ nghĩa tư bản hiện đại", "84", "34 / 34 / 16"),
        ("5", "Kinh tế thị trường định hướng XHCN và quan hệ lợi ích", "84", "34 / 34 / 16"),
        ("6", "Công nghiệp hóa, cách mạng công nghiệp và hội nhập", "84", "34 / 34 / 16"),
    ]
    add_data_table(
        doc,
        ["Chương", "Trọng tâm", "Số câu", "NB / TH / VD"],
        question_rows,
        [1080, 5080, 1040, 2160],
        center_columns=(0, 2, 3),
        total_row=("Tổng", "Đủ 6 chương", "504", "204 / 204 / 96"),
    )
    add_body(
        doc,
        "Ngân hàng gồm 204 câu Nhận biết, 204 câu Thông hiểu và 96 câu Vận dụng. Mỗi câu có bốn lựa chọn A-D, một đáp án đúng, phần giải thích, thông tin chương/slot/chủ đề và dẫn trang tài liệu. Phân bố đáp án được cân bằng theo từng chương để hạn chế việc đoán dựa trên vị trí.",
    )

    add_heading(doc, "2.3. Kiểm định nội dung", 2)
    add_list_item(doc, bullet_num, " Dữ liệu được tách theo sáu tệp chương, sau đó hợp nhất thành ngân hàng sản xuất bằng pipeline tự động.")
    add_list_item(doc, bullet_num, " Validator kiểm tra schema, ID, nguồn, trang, đáp án, nội dung HTML rác, trùng lặp và dấu hiệu thiên lệch độ dài phương án.")
    add_list_item(doc, bullet_num, " Đợt mở rộng bổ sung 204 câu mới, chia đều 34 câu cho mỗi chương và dùng 141 trang PDF khác nhau.")
    add_list_item(doc, bullet_num, " Kết quả kiểm định gần nhất: 504 câu hợp lệ, 0 lỗi và 0 cảnh báo.")

    add_heading(doc, "3. Các chức năng học tập", 1)
    add_heading(doc, "3.1. Chế độ Luyện thi", 2)
    add_body(
        doc,
        "Luyện thi là chế độ làm câu hỏi trắc nghiệm có chấm điểm. Người học có thể lọc theo chương và mức độ, xáo trộn thứ tự, chỉ xem câu đã đánh dấu hoặc tập trung vào nhóm hay trả lời sai. Thanh thống kê cho biết vị trí hiện tại, điểm của phiên, số câu đã luyện và số câu đã đánh dấu.",
    )
    add_list_item(doc, bullet_num, " Chọn một phương án để nhận trạng thái đúng/sai ngay lập tức.")
    add_list_item(doc, bullet_num, " Sau khi trả lời, hệ thống hiện đáp án đúng, giải thích và nguồn tài liệu.")
    add_list_item(doc, bullet_num, " Có thể quay lại câu trước, chuyển câu sau và xem lại lựa chọn đã trả lời.")
    add_list_item(doc, bullet_num, " Nút đánh dấu giúp tạo danh sách ôn tập cá nhân; bộ lọc “Hay sai” tập trung vào kiến thức còn yếu.")

    add_heading(doc, "3.2. Chế độ Flashcard", 2)
    add_body(
        doc,
        "Flashcard dùng cùng bộ lọc và thứ tự câu hỏi nhưng không buộc người học phải chọn đáp án để chấm điểm. Thao tác lật thẻ hiển thị đáp án, giải thích và nguồn. Cách này phù hợp với ôn nhanh, tự trả lời trong đầu hoặc học theo nhóm.",
    )
    add_figure(
        doc,
        FLASHCARD_SHOT,
        "Hình 2. Flashcard sau khi lật thẻ, hiển thị đáp án, giải thích và nguồn.",
        width=6.18,
        alt_text="Giao diện Flashcard đã lật, làm nổi bật đáp án đúng và phần giải thích.",
    )

    add_heading(doc, "3.3. Chế độ Tìm kiếm", 2)
    add_body(
        doc,
        "Tìm kiếm cho phép tra cứu toàn bộ 504 câu theo từ khóa. Hệ thống hỗ trợ nhập không dấu, làm nổi bật phần khớp và nhóm kết quả theo chương. Chế độ này hữu ích khi cần ôn một khái niệm cụ thể như giá trị thặng dư, tư bản tài chính, quan hệ lợi ích hoặc hội nhập kinh tế quốc tế.",
    )

    add_heading(doc, "3.4. Đánh dấu, thống kê và lưu tiến độ", 2)
    add_body(
        doc,
        "Website tự động lưu tiến độ học trên trình duyệt. Phiên Luyện thi và Flashcard được lưu độc lập, gồm bộ lọc, thứ tự câu sau khi xáo trộn, vị trí hiện tại, các lựa chọn đã trả lời, điểm phiên và trạng thái lật thẻ. Dữ liệu được cập nhật sau mỗi câu trả lời, khi chuyển câu, đổi chế độ, tải lại trang hoặc tạm rời trang.",
    )
    add_callout(
        doc,
        "Lợi ích",
        "Người học có thể làm đến câu thứ 10, chuyển sang mục khác hoặc mở trang web khác rồi quay lại mà không mất vị trí và kết quả. Nút Đặt lại xóa tiến độ học khi người dùng muốn bắt đầu một phiên mới.",
    )

    add_heading(doc, "4. Chế độ game Công thành", 1, page_break_before=True)
    add_heading(doc, "4.1. Mục tiêu và khởi tạo chiến dịch", 2)
    add_body(
        doc,
        "Công thành là game chiến lược theo lượt, trong đó kiến thức MLN222 quyết định một phần năng lực phát triển thế lực. Người chơi chọn một trong 34 tỉnh/thành làm căn cứ khởi đầu; 33 địa bàn còn lại do các thế lực NPC cát cứ. Mỗi tỉnh có vùng, địa hình, quy mô sức chứa và đặc tính kinh tế hoặc phòng thủ riêng.",
    )
    add_body(
        doc,
        "Mục tiêu chiến dịch là mở rộng quyền kiểm soát trên bản đồ, quản lý dân số và tài nguyên, xây dựng quân đội, thiết lập quan hệ với các thế lực khác và hoàn thành điều kiện thống nhất. Chiến dịch không có giới hạn lượt cứng nên người chơi có thể tiếp tục từ vị trí khởi đầu khó.",
    )

    add_heading(doc, "4.2. Bản đồ và điều khiển", 2)
    add_list_item(doc, bullet_num, " Bản đồ Việt Nam gồm 34 tỉnh/thành và các nhóm đảo, được nhúng trực tiếp để chạy ngoại tuyến.")
    add_list_item(doc, bullet_num, " Hỗ trợ phóng to, thu nhỏ, vừa khung, tập trung tỉnh, kéo bản đồ, con lăn chuột và pinch trên màn hình cảm ứng.")
    add_list_item(doc, bullet_num, " Nhấp chuột phải vào tỉnh của thế lực khác mở vòng lệnh Thông tin, Ngoại giao, Thương mại và Quân sự.")
    add_list_item(doc, bullet_num, " Bàn phím dùng phím Context Menu hoặc Shift+F10; thiết bị cảm ứng dùng nút Hành động hoặc nhấn giữ.")
    add_figure(
        doc,
        WHEEL_SHOT,
        "Hình 3. Vòng lệnh theo ngữ cảnh khi thao tác với một thế lực khác.",
        width=6.18,
        alt_text="Giao diện vòng lệnh gồm Thông tin, Ngoại giao, Thương mại và Quân sự trên bản đồ.",
    )

    add_heading(doc, "4.3. Chu kỳ lượt và điểm lệnh", 2)
    add_body(
        doc,
        "Mỗi lượt người chơi có 2 điểm lệnh. Các thao tác hợp lệ được xếp vào khay lệnh và có thể hủy riêng trước khi kết thúc lượt. Khi xác nhận, hệ thống giải quyết kinh tế, quyết định của NPC, giao tranh, chiếm đóng và phần thưởng quiz theo một thứ tự xác định.",
    )
    add_list_item(doc, bullet_num, " Tuyển quân hoặc mở khóa binh chủng mới.")
    add_list_item(doc, bullet_num, " Di chuyển và viện binh giữa các tỉnh do người chơi kiểm soát.")
    add_list_item(doc, bullet_num, " Mở tuyến thương mại, đề nghị hiệp ước hoặc phản hồi đề nghị ngoại giao.")
    add_list_item(doc, bullet_num, " Cảnh báo tấn công, chuẩn bị chiến dịch quân sự hoặc chờ để bảo toàn nguồn lực.")

    add_heading(doc, "4.4. Dân số, kinh tế và duy trì", 2)
    add_body(
        doc,
        "Tài nguyên chính gồm lương thực, tiền, dân thường, quân đội và giới hạn dân số. Dân thường tạo ra lương thực và tiền; quân đội tiêu tốn dân thường khi tuyển và cần chi phí duy trì sau mỗi lượt. Vì vậy, huy động quá nhiều quân làm giảm lực lượng lao động, sản xuất và tốc độ tăng dân số.",
    )
    add_body(
        doc,
        "Tăng trưởng dân thường đạt mức cao nhất khi số dân thường xấp xỉ 40% sức chứa. Dưới ngưỡng này, tốc độ tăng dần; vượt ngưỡng, tốc độ giảm dần về 0 khi gần đầy. Tỷ lệ quân đội cao tiếp tục làm giảm tăng trưởng và năng suất. Nếu thiếu lương thực hoặc tiền kéo dài, quân đội bị giảm tiếp tế, tinh thần và có thể đào ngũ.",
    )
    add_callout(
        doc,
        "Cân bằng chiến lược",
        "Mở rộng lãnh thổ làm tăng sức chứa và tiềm lực dài hạn, nhưng tỉnh mới chiếm chịu ba lượt ổn định với sản xuất và tăng trưởng bị giảm. Người chơi phải cân đối giữa mở rộng nhanh, bảo vệ hậu phương và duy trì nền kinh tế.",
        fill=PALE_GOLD,
        accent="B47000",
    )

    add_heading(doc, "4.5. Binh chủng và mở khóa", 2)
    unit_rows = [
        ("Dân binh", "Phòng thủ lãnh thổ, đồn trú", "Có sẵn từ đầu"),
        ("Bộ binh", "Lực lượng chiến tuyến", "Từ lượt 3, cần tiền và Dân binh"),
        ("Xạ binh", "Hỗ trợ tầm xa", "Từ lượt 5, cần 2 tỉnh và Bộ binh"),
        ("Kỵ binh", "Cơ động, đột kích", "Từ lượt 8, cần 3 tỉnh và Bộ binh"),
        ("Công binh", "Hỗ trợ công thành", "Từ lượt 8, cần 3 tỉnh và Bộ binh"),
    ]
    add_data_table(
        doc,
        ["Binh chủng", "Vai trò", "Điều kiện khái quát"],
        unit_rows,
        [1900, 3000, 4460],
        center_columns=(0,),
    )
    add_body(
        doc,
        "Mỗi binh chủng có sức mạnh, chi phí tuyển và mức duy trì khác nhau. Lực lượng cơ động hoặc chuyên môn cao tạo lợi thế nhưng đòi hỏi nền kinh tế đủ lớn. Số quân luôn nằm trong giới hạn dân số chung của các tỉnh.",
    )

    add_heading(doc, "4.6. Ngoại giao và thương mại", 2)
    add_body(
        doc,
        "Quan hệ giữa các thế lực được biểu diễn trên thang từ đối địch đến thân thiện. Người chơi có thể mở tuyến thương mại để nhận tiền mỗi lượt, đề nghị hiệp ước không xâm phạm hoặc liên minh khi quan hệ đạt điều kiện. Số tuyến và hiệp ước bị giới hạn để tránh tích lũy lợi ích không kiểm soát; hành vi phản bội làm giảm mạnh quan hệ và khóa khả năng ngoại giao trong một thời gian.",
    )

    add_heading(doc, "4.7. Chiến đấu nhiều lượt", 2)
    add_body(
        doc,
        "Tấn công không tạo kết quả ngay lập tức. Thế lực tiến công phải cảnh báo trước một lượt; sau đó trận chiến được giải quyết qua nhiều nhịp dựa trên quân số, thành phần binh chủng, địa hình, công sự, tinh thần và tiếp tế. Người chơi có thể chọn chiến thuật công thành, giao chiến, tổng công kích, củng cố hoặc rút lui, đồng thời điều viện binh từ hậu phương.",
    )
    add_list_item(doc, bullet_num, " Công thành giảm thương vong nhưng cần thời gian để phá công sự.")
    add_list_item(doc, bullet_num, " Tổng công kích gây sát thương lớn hơn nhưng làm quân tấn công chịu rủi ro cao.")
    add_list_item(doc, bullet_num, " Củng cố hồi tinh thần, còn rút lui bảo toàn một phần lực lượng nhưng có thể bị truy kích.")
    add_list_item(doc, bullet_num, " Thương vong được tách thành tử trận, bị thương và tan rã; quân bị thương cần thời gian hồi phục.")
    add_body(
        doc,
        "Một tỉnh chỉ đổi chủ khi lực lượng phòng thủ bị đánh bại hoặc tan rã. Sau chiến thắng, tỉnh bước vào giai đoạn chiếm đóng ba lượt trước khi hoạt động bình thường.",
    )

    add_heading(doc, "4.8. NPC và điều kiện chiến thắng", 2)
    add_body(
        doc,
        "NPC có bốn thiên hướng chính: thận trọng, thương nhân, bành trướng và phòng thủ. Mỗi thế lực tự cân nhắc phát triển kinh tế, tuyển quân, ngoại giao, tăng viện hoặc tấn công; người chơi cũng có thể bị đánh bất ngờ sau giai đoạn đầu chiến dịch.",
    )
    add_body(
        doc,
        "Chiến thắng yêu cầu kiểm soát tối thiểu 60% điểm lãnh thổ toàn quốc và ít nhất 4 trong 6 vùng. Điểm lãnh thổ được tính theo quy mô tỉnh, vì vậy kiểm soát nhiều tỉnh nhỏ chưa chắc tương đương với giữ các địa bàn lớn. Người chơi thất bại khi không còn tỉnh nào.",
    )

    add_heading(doc, "4.9. Thử thách 10 câu cuối lượt", 2)
    add_body(
        doc,
        "Kết thúc mỗi lượt, người chơi phải hoàn thành 10 câu hỏi ngẫu nhiên, không trùng trong cùng bài. Mỗi câu hiển thị phản hồi, đáp án đúng, giải thích và nguồn. Người chơi được chọn ưu tiên phần thưởng: lương thực, tiền, dân số hoặc mở khóa.",
    )
    reward_rows = [
        ("0-2", "Phạt nặng", "Trừ một phần tài nguyên và giảm sản xuất trong lượt sau"),
        ("3-4", "Phạt nhẹ", "Khấu trừ nhỏ hơn; không cộng thưởng"),
        ("5", "Trung hòa", "Không thưởng, không phạt"),
        ("6-7", "Thưởng nhỏ", "Cộng tài nguyên hoặc dân nhập cư theo ưu tiên"),
        ("8-9", "Thưởng lớn", "Cộng tài nguyên và tăng sản xuất trong 2 lượt"),
        ("10", "Hoàn hảo", "Thưởng tối đa; có thể nhận giảm giá mở khóa trong 2 lượt"),
    ]
    add_data_table(
        doc,
        ["Điểm", "Mức", "Tác động chính"],
        reward_rows,
        [950, 1800, 6610],
        center_columns=(0, 1),
    )
    add_figure(
        doc,
        QUIZ_SHOT,
        "Hình 4. Quiz 10 câu cuối lượt trên thiết bị di động, có phản hồi và giải thích.",
        width=2.42,
        alt_text="Ảnh giao diện quiz cuối lượt trên điện thoại, hiển thị đáp án đúng và giải thích.",
    )

    add_heading(doc, "4.10. Lưu và tiếp tục chiến dịch", 2)
    add_body(
        doc,
        "Chiến dịch được lưu sau mỗi lệnh hợp lệ, thay đổi chiến thuật, câu trả lời quiz và khi hoàn tất lượt. Khi mở lại website, nút Tiếp tục khôi phục cả trạng thái chiến dịch và các lệnh đang chờ. Bản lưu được kiểm tra phiên bản, kích thước, checksum và các bất biến của game; bản lưu hỏng không làm ứng dụng ngừng hoạt động.",
    )

    add_heading(doc, "5. Thiết kế giao diện và khả năng truy cập", 1, page_break_before=True)
    add_heading(doc, "5.1. Hệ thống giao diện lai", 2)
    add_body(
        doc,
        "Không gian học sử dụng nền sáng, khoảng trắng rõ và màu xanh ngọc làm điểm nhấn. Không gian game chuyển sang nền tối chiến thuật để bản đồ, tài nguyên và cảnh báo nổi bật hơn. Dù khác sắc thái, hai phần vẫn dùng chung hệ thống biểu tượng, nhịp khoảng cách và cách tổ chức điều khiển.",
    )
    add_heading(doc, "5.2. Tương thích thiết bị", 2)
    add_list_item(doc, bullet_num, " Desktop dùng thanh điều hướng trái, bản đồ lớn và bảng điều hành bên phải.")
    add_list_item(doc, bullet_num, " Mobile dùng điều hướng đáy, HUD tài nguyên thu gọn và bảng chiến dịch dạng bottom sheet.")
    add_list_item(doc, bullet_num, " Hỗ trợ màn hình dọc, ngang, safe area và tình huống bàn phím ảo xuất hiện.")
    add_list_item(doc, bullet_num, " Không có cuộn ngang tại các viewport kiểm thử 1440x900, 1024x768, 390x844, 360x800 và 844x390.")

    add_heading(doc, "5.3. Khả năng truy cập", 2)
    add_list_item(doc, bullet_num, " Các nút tương tác chính có kích thước tối thiểu 44 px trên thiết bị nhỏ.")
    add_list_item(doc, bullet_num, " Trạng thái đúng, sai, cảnh báo và khóa thao tác dùng đồng thời màu, biểu tượng và nội dung chữ.")
    add_list_item(doc, bullet_num, " Bản đồ và vòng lệnh có thể thao tác bằng bàn phím; focus được khôi phục khi đóng menu.")
    add_list_item(doc, bullet_num, " Modal quiz giữ focus trong phạm vi cần thiết và không cho đóng bằng Escape khi chưa hoàn tất.")
    add_list_item(doc, bullet_num, " Giao diện tôn trọng thiết lập prefers-reduced-motion và có độ tương phản đạt ngưỡng kiểm thử.")

    add_heading(doc, "6. Thiết kế kỹ thuật và triển khai", 1)
    add_heading(doc, "6.1. Kiến trúc dữ liệu và mã nguồn", 2)
    add_body(
        doc,
        "Mã nguồn được tách theo trách nhiệm: dữ liệu câu hỏi theo chương; dữ liệu bản đồ và cân bằng game; engine kinh tế, dân số, ngoại giao, chiến đấu và AI; lớp lưu trữ; lớp điều khiển; và các mô-đun giao diện. Trạng thái game là JSON thuần, giúp kiểm tra, lưu, khôi phục và mô phỏng không phụ thuộc DOM.",
    )
    add_body(
        doc,
        "Engine sử dụng các luồng số ngẫu nhiên xác định cho AI, chiến đấu, sự kiện và quiz. Với cùng seed, trạng thái ban đầu và chuỗi hành động, kết quả đầu ra tương đương theo byte. Đặc tính này hỗ trợ tái hiện lỗi và kiểm thử cân bằng.",
    )

    add_heading(doc, "6.2. Đóng gói ngoại tuyến", 2)
    add_body(
        doc,
        "Pipeline hợp nhất 504 câu, kiểm tra dữ liệu game, nén tài nguyên bản đồ, nhúng CSS/JavaScript/ảnh và sinh index.html. File đầu ra có thể mở trực tiếp bằng giao thức file://, không cần cài máy chủ và không gửi yêu cầu HTTP(S) khi chạy.",
    )
    add_callout(
        doc,
        "Phân tách dữ liệu lưu",
        "Tiến độ học, chiến dịch và trạng thái giao diện game dùng các khóa lưu riêng. Việc đặt lại game không xóa dữ liệu học; đặt lại phiên học không làm mất chiến dịch.",
    )

    add_heading(doc, "6.3. An toàn dữ liệu cục bộ", 2)
    add_body(
        doc,
        "Ứng dụng xử lý trường hợp localStorage không khả dụng, vượt dung lượng hoặc chứa JSON hỏng. Dữ liệu không hợp lệ bị từ chối có kiểm soát, trong khi phiên đang chạy trong bộ nhớ vẫn được giữ nếu có thể. Nội dung câu hỏi và kết quả tìm kiếm được dựng bằng DOM node thay vì chèn HTML động từ dữ liệu.",
    )

    add_heading(doc, "7. Kiểm thử và chất lượng", 1)
    add_body(
        doc,
        "Hệ thống được kiểm thử ở ba lớp: chất lượng nội dung, logic dữ liệu/engine và hành vi trình duyệt. Bảng dưới phản ánh trạng thái kiểm định tại thời điểm lập báo cáo ngày 19/07/2026.",
    )
    qa_rows = [
        ("Ngân hàng câu hỏi", "504 câu; 6 chương", "PASS - 0 lỗi, 0 cảnh báo"),
        ("Dữ liệu game", "34 tỉnh; 6 vùng; 58 cạnh; 5 binh chủng", "PASS"),
        ("Python pipeline", "36 bài kiểm thử", "PASS 36/36"),
        ("Node engine/UI", "139 bài kiểm thử", "PASS 139/139"),
        ("Mô phỏng kinh tế", "100.000 bước chuyển", "PASS bất biến"),
        ("Mô phỏng chiến đấu", "10.000 trận", "PASS bất biến"),
        ("Mô phỏng chiến dịch", "1.000 chiến dịch x 60 lượt", "PASS gate cân bằng"),
        ("Trình duyệt", "Desktop, mobile, offline, save/resume", "PASS"),
    ]
    add_data_table(
        doc,
        ["Hạng mục", "Phạm vi", "Kết quả"],
        qa_rows,
        [2450, 4050, 2860],
        center_columns=(2,),
    )
    add_body(
        doc,
        "Bản build hiện tại có kích thước 1.801.389 byte và hai lần build liên tiếp cho cùng SHA-256 205653B3480BE808FB0AAFE2C82E5F8B7F7BFE5DCD86D53BC21A4EE5C18CD893. Kiểm thử trình duyệt xác nhận không có lỗi console, page error hoặc yêu cầu mạng ngoài trong workflow từ học tập đến hoàn tất quiz và sang lượt 2.",
    )

    add_heading(doc, "8. Hướng dẫn sử dụng nhanh", 1, page_break_before=True)
    add_list_item(doc, decimal_num, "Mở tệp index.html bằng Chrome, Edge hoặc trình duyệt hiện đại.")
    add_list_item(doc, decimal_num, "Chọn Luyện thi để làm câu hỏi có chấm điểm, Flashcard để ôn nhanh hoặc Tìm kiếm để tra cứu khái niệm.")
    add_list_item(doc, decimal_num, "Thiết lập chương, mức độ và tùy chọn xáo trộn; trả lời câu hỏi rồi đọc phần giải thích và nguồn.")
    add_list_item(doc, decimal_num, "Chọn Công thành, chọn tỉnh khởi đầu và mã chiến dịch, sau đó bắt đầu lượt 1.")
    add_list_item(doc, decimal_num, "Theo dõi tài nguyên, chọn tỉnh trên bản đồ và dùng bảng hành động hoặc chuột phải để xếp lệnh.")
    add_list_item(doc, decimal_num, "Chọn ưu tiên thưởng, kết thúc lượt và hoàn thành đủ 10 câu hỏi để nhận kết quả.")
    add_list_item(doc, decimal_num, "Dùng nút Lưu hoặc để cơ chế tự lưu hoạt động; lần sau chọn Tiếp tục để khôi phục chiến dịch.")
    add_callout(
        doc,
        "Lưu ý",
        "Tiến độ được lưu trong dữ liệu cục bộ của trình duyệt. Việc xóa dữ liệu duyệt web, dùng chế độ ẩn danh, đổi trình duyệt hoặc môi trường file:// khác có thể làm bản lưu không còn truy cập được.",
        fill=PALE_GOLD,
        accent="B47000",
    )

    add_heading(doc, "9. Giá trị sử dụng và hướng phát triển", 1)
    add_heading(doc, "9.1. Giá trị sử dụng", 2)
    add_list_item(doc, bullet_num, " Chuyển việc đọc thụ động thành luyện truy hồi kiến thức có phản hồi.")
    add_list_item(doc, bullet_num, " Bao phủ nhiều lớp lý thuyết, từ khái niệm đến quan hệ và tình huống vận dụng.")
    add_list_item(doc, bullet_num, " Tạo động lực quay lại học qua chiến dịch dài hạn và phần thưởng theo kết quả quiz.")
    add_list_item(doc, bullet_num, " Cho phép học ngoại tuyến, không cần tài khoản hoặc cài đặt máy chủ.")
    add_list_item(doc, bullet_num, " Cung cấp nguồn dẫn để người học kiểm tra lại thay vì chỉ ghi nhớ đáp án.")

    add_heading(doc, "9.2. Giới hạn hiện tại", 2)
    add_list_item(doc, bullet_num, " Dữ liệu chỉ lưu cục bộ, chưa đồng bộ giữa nhiều thiết bị và chưa có tài khoản người dùng.")
    add_list_item(doc, bullet_num, " Độ khó chiến dịch thay đổi theo tỉnh khởi đầu; một số vị trí cần nhiều lượt hơn để đạt điều kiện thắng.")
    add_list_item(doc, bullet_num, " Kiểm thử giao diện có ảnh minh chứng và kiểm tra hình học DOM nhưng chưa có pixel-diff tự động trong CI.")

    add_heading(doc, "9.3. Hướng phát triển đề xuất", 2)
    add_list_item(doc, bullet_num, " Đồng bộ tiến độ qua tài khoản và hỗ trợ xuất/nhập bản lưu chiến dịch.")
    add_list_item(doc, bullet_num, " Bổ sung thống kê theo chủ đề, đường cong ghi nhớ và gợi ý ôn tập thích ứng.")
    add_list_item(doc, bullet_num, " Thêm hướng dẫn nhập môn, mức khó chiến dịch và sự kiện lịch sử/kinh tế có kiểm soát.")
    add_list_item(doc, bullet_num, " Mở rộng kiểm thử trực quan tự động và đo khả năng truy cập trên thiết bị thật.")

    add_heading(doc, "10. Kết luận", 1)
    add_body(
        doc,
        "Website MLN222 đã hình thành một môi trường tự học hoàn chỉnh gồm ngân hàng 504 câu có nguồn, ba công cụ học tập chuyên biệt và game Công thành có chiều sâu về kinh tế, dân số, ngoại giao và chiến đấu. Sản phẩm vừa đáp ứng nhu cầu ôn thi nhanh, vừa tạo một trải nghiệm học qua chơi có thể kéo dài qua nhiều phiên.",
    )
    add_body(
        doc,
        "Khả năng chạy ngoại tuyến, lưu tiến độ, giao diện đáp ứng nhiều thiết bị và hệ thống kiểm thử tự động giúp website phù hợp để sử dụng như công cụ học tập cá nhân hoặc tài liệu bổ trợ cho học phần MLN222.",
    )

    add_heading(doc, "Tài liệu tham khảo", 1)
    add_source_entry(
        doc,
        source_num,
        "Giáo trình chuẩn: ",
        r"F:\MLN222\GIAO-TRINH-KINH-TE-CHINH-TRI-MAC-LENIN-BO-GIAO-DUC-VA-DAO-TAO.pdf",
    )
    add_source_entry(doc, source_num, "Bài giảng bổ trợ: ", r"Các tệp *.pptx.txt theo slot trong F:\MLN222.")
    add_source_entry(
        doc,
        source_num,
        "Nguồn hình học bản đồ: ",
        "Bản đồ Việt Nam 34 tỉnh/thành - FPT Lịch sử Việt",
        url="https://fptlichsuviet.io.vn/map/index.html",
    )
    add_source_entry(doc, source_num, "Mã nguồn và dữ liệu sản phẩm: ", str(ROOT))
    add_source_entry(
        doc,
        source_num,
        "Báo cáo kiểm thử: ",
        r"plans\260719-1752-hybrid-ui-command-wheel\reports\end-to-end-testing.md",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_report()
    print(f"Created {path}")
